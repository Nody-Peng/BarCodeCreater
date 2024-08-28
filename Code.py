import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox
from tkinter import Listbox, Scrollbar, Entry, Label, Button
from tkinter import SINGLE
from barcode import Code128
from barcode.writer import ImageWriter
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches
import os

# 添加產品到列表
def add_product():
    name = name_entry.get()
    code = code_entry.get()
    if name and code:
        products.append((code, name))
        products_listbox.insert(tk.END, f'{name} - {code}')
        name_entry.delete(0, tk.END)
        code_entry.delete(0, tk.END)
    else:
        messagebox.showwarning("輸入錯誤", "請填寫產品名稱和條碼。")

# 刪除選定的產品
def delete_product():
    selected_index = products_listbox.curselection()
    if selected_index:
        products_listbox.delete(selected_index)
        del products[selected_index[0]]
    else:
        messagebox.showwarning("選擇錯誤", "請選擇要刪除的產品。")

# 選擇條碼輸出目錄
def select_output_folder():
    global barcode_output_folder
    barcode_output_folder = filedialog.askdirectory()
    output_folder_entry.delete(0, tk.END)
    output_folder_entry.insert(0, barcode_output_folder)

# 選擇 Word 文件輸出目錄
def select_output_word_file():
    global word_output_file
    word_output_file = filedialog.asksaveasfilename(defaultextension=".docx",
                                                   filetypes=[("Word 文件", "*.docx")])
    output_word_file_entry.delete(0, tk.END)
    output_word_file_entry.insert(0, word_output_file)

# 生成 Word 文件
def generate():
    if not products:
        messagebox.showwarning("列表為空", "請添加產品到列表。")
        return

    if not barcode_output_folder or not word_output_file:
        messagebox.showwarning("路徑未選", "請選擇條碼和 Word 文件的輸出目錄。")
        return

    def generate_barcode_with_text(code, name):
        options = {
            'module_width': 0.5,
            'module_height': 15.0,
            'font_size': 10,
            'text_distance': 5,
            'quiet_zone': 6.5,
        }
        barcode_class = Code128(code, writer=ImageWriter())
        barcode_path = os.path.join(barcode_output_folder, f'{name}.png')
        barcode_class.save(os.path.join(barcode_output_folder, name), options=options)

        barcode_img = Image.open(barcode_path)
        large_font_size = 50
        try:
            # 嘗試家載支持中文的字體
            font_large = ImageFont.truetype("msyh.ttc", large_font_size)
        except IOError:
            font_large = ImageFont.load_default()

        draw = ImageDraw.Draw(barcode_img)
        name_bbox = draw.textbbox((0, 0), name, font=font_large)
        name_width = name_bbox[2] - name_bbox[0]
        name_height = name_bbox[3] - name_bbox[1]

        # 調整條碼和物品名稱之間的間距
        vertical_gap = 40  # 設置間隔

        total_height = barcode_img.height + name_height + vertical_gap
        total_width = max(barcode_img.width, name_width)

        result_img = Image.new('RGB', (total_width, total_height), 'white')
        draw = ImageDraw.Draw(result_img)
        result_img.paste(barcode_img, ((total_width - barcode_img.width) // 2, name_height + vertical_gap))
        draw.text(((total_width - name_width) // 2, 0), name, fill='black', font=font_large)

        result_img_path = os.path.join(barcode_output_folder, f'{name}_with_text.png')
        result_img.save(result_img_path)
        return result_img_path

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)

    for code, name in products:
        barcode_img_path = generate_barcode_with_text(code, name)
        doc.add_picture(barcode_img_path, width=Inches(2))
        doc.add_paragraph()

    doc.save(word_output_file)
    messagebox.showinfo("成功", f'Word 文件已生成並保存為 {word_output_file}')

# 創建主窗口
root = tk.Tk()
root.title("條碼生成器")
root.geometry("600x500")

products = []
barcode_output_folder = ""
word_output_file = ""

# GUI設計
Label(root, font=10, text="產品名稱:").grid(row=0, column=0, padx=10, pady=10)
name_entry = Entry(root, font=10, width=30)
name_entry.grid(row=0, column=1, padx=10, pady=10)

Label(root, font=10, text="產品條碼:").grid(row=1, column=0, padx=10, pady=10)
code_entry = Entry(root, font=10, width=30)
code_entry.grid(row=1, column=1, padx=10, pady=10)

Button(root, font=10, text="添加產品", command=add_product).grid(row=0, column=2, padx=10, pady=10)

# 調整 Listbox 的位置和大小
products_listbox = Listbox(root, font=10, selectmode=SINGLE, width=55, height=12)
products_listbox.grid(row=2, column=0, columnspan=2, padx=10, pady=10)

scrollbar = Scrollbar(root)
scrollbar.grid(row=2, column=2, sticky='ns', padx=(0, 10))
products_listbox.config(yscrollcommand=scrollbar.set)
scrollbar.config(command=products_listbox.yview)

# 調整刪除按鈕的位置
Button(root, font=10, text="刪除選定項目", command=delete_product).grid(row=1, column=2, columnspan=3, padx=10, pady=10)

Label(root, font=10, text="選擇條碼輸出位置:").grid(row=4, column=0, padx=10, pady=10)
output_folder_entry = Entry(root, font=10, width=30)
output_folder_entry.grid(row=4, column=1, padx=10, pady=10)
Button(root, font=10, text="瀏覽", command=select_output_folder).grid(row=4, column=2, padx=10, pady=10)

Label(root, font=10, text="選擇 Word 文件輸出位置:").grid(row=5, column=0, padx=10, pady=10)
output_word_file_entry = Entry(root, font=10, width=30)
output_word_file_entry.grid(row=5, column=1, padx=10, pady=10)
Button(root, font=10, text="瀏覽", command=select_output_word_file).grid(row=5, column=2, padx=10, pady=10)

Button(root, font=10, text="生成 Word 文件", command=generate).grid(row=6, column=1, pady=20)

root.mainloop()
